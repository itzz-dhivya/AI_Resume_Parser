from dotenv import load_dotenv
load_dotenv()
import streamlit as st
import os
import hashlib
import pandas as pd
import mysql.connector
from mysql.connector import Error
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
import tempfile
import mysql.connector
from mysql.connector import Error
import requests
import time
from streamlit_lottie import st_lottie
import plotly.express as px
import base64
import re
from docx import Document
import tempfile
import os
from io import BytesIO

# ----------------- CONFIG -----------------
st.set_page_config(page_title="AI Resume Analyzer", layout="centered", initial_sidebar_state="auto")

# ----------------- DATABASE -----------------#

import streamlit as st
from config.db_config import get_conn

# ----------------- USER LOGIN EXAMPLE -----------------
def check_user(username, password):
    """Validate login credentials"""
    conn = get_conn()
    if conn:
        try:
            cursor = conn.cursor(dictionary=True)
            query = "SELECT * FROM users WHERE username=%s AND password=%s"
            cursor.execute(query, (username, password))
            user = cursor.fetchone()
            return user
        except Exception as e:
            st.error(f"DB Query Error: {e}")
            return None
        finally:
            cursor.close()
            conn.close()
    return None

# ----------------- USER REGISTRATION EXAMPLE -----------------
def register_user(username, password, email):
    """Register a new user"""
    conn = get_conn()
    if conn:
        try:
            cursor = conn.cursor()
            query = "INSERT INTO users (username, password, email) VALUES (%s, %s, %s)"
            cursor.execute(query, (username, password, email))
            conn.commit()
            st.success("✅ User registered successfully")
        except Exception as e:
            st.error(f"DB Insert Error: {e}")
        finally:
            cursor.close()
            conn.close()

# ----------------- USAGE IN APP -----------------
def login_module():
    st.title("Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        user = check_user(username, password)
        if user:
            st.success(f"Welcome {user['username']}!")
        else:
            st.error("Invalid credentials")



# ----------------- HASHING -----------------
SALT = "a_random_salt_123!@#"
def hash_password(password: str) -> str:
    return hashlib.sha256((SALT + password).encode()).hexdigest()

def check_password(password: str, hashed: str) -> bool:
    return hash_password(password) == hashed

# ----------------- STYLES -----------------
def set_bg_and_styles():
    image_path = "assets/bg/analy.jpg"
    if os.path.exists(image_path):
        with open(image_path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode()
        st.markdown(f"""
        <style>
        .stApp {{
            background-image: url("data:image/jpg;base64,{encoded}");
            background-size: cover;
            background-position: center;
            background-attachment: fixed;
            color: #e6eef8;
        }}
        .glass-card {{
            background: rgba(255, 255, 255, 0.06);
            backdrop-filter: blur(6px);
            border-radius: 12px;
            padding: 18px;
            box-shadow: 0 8px 30px rgba(2,6,23,0.6);
            border: 1px solid rgba(255,255,255,0.06);
            margin-bottom: 16px;
        }}
        [data-testid="stSidebar"] {{
            background: rgba(255,255,255,0.05);
            backdrop-filter: blur(8px);
            border-radius: 12px;
            padding: 10px;
            border: 1px solid rgba(255,255,255,0.1);
        }}
        .small-note {{ font-size:12px; color:#c9d7f0; }}
        </style>
        """, unsafe_allow_html=True)

set_bg_and_styles()

def card_container(content):
    st.markdown(f"<div class='glass-card'>{content}</div>", unsafe_allow_html=True)

def load_lottie(url):
    try:
        r = requests.get(url, timeout=6)
        if r.status_code == 200:
            return r.json()
    except:
        return None
    return None

# ----------------- ROLE SKILLS & COURSES -----------------
# ----------------- ROLE SKILLS & COURSES -----------------
role_skills_courses ={

    # ---------------- Software / Web / Mobile Developers ----------------
    "Software Developer": {
        "skills": ["Python", "Java", "C++", "C#", "SQL", "Git"],
        "courses": {
            "Python": [
                "https://www.coursera.org/learn/python",
                "https://learn.microsoft.com/en-us/training/paths/python-language/",
                "https://www.udemy.com/course/pythonforbeginnersfree/"
            ],
            "Java": [
                "https://www.coursera.org/learn/java-programming",
                "https://learn.microsoft.com/en-us/training/paths/java-fundamentals/"
            ],
            "C++": [
                "https://www.coursera.org/specializations/c-plus-plus-modern",
                "https://www.udemy.com/course/free-learn-cpp-tutorial-course/"
            ],
            "C#": [
                "https://learn.microsoft.com/en-us/training/paths/csharp-first-steps/",
                "https://www.coursera.org/learn/c-sharp"
            ],
            "SQL": [
                "https://www.coursera.org/learn/sql-for-data-science",
                "https://learn.microsoft.com/en-us/training/paths/query-relational-data-sql/"
            ],
            "Git": [
                "https://www.coursera.org/learn/introduction-git-github",
                "https://www.udemy.com/course/git-tutorial-free/"
            ]
        }
    },

    "Front-End Developer": {
        "skills": ["HTML", "CSS", "JavaScript", "React", "Vue.js"],
        "courses": {
            "HTML": [
                "https://www.coursera.org/learn/html-css-javascript",
                "https://www.udemy.com/course/html-tutorial-free/"
            ],
            "CSS": [
                "https://www.coursera.org/learn/html-css-javascript",
                "https://www.udemy.com/course/css-tutorial-free/"
            ],
            "JavaScript": [
                "https://www.coursera.org/learn/javascript-basics",
                "https://www.udemy.com/course/javascript-tutorial-free/"
            ],
            "React": [
                "https://www.coursera.org/learn/react-basics",
                "https://www.udemy.com/course/learn-react-js-tutorial-free/"
            ],
            "Vue.js": [
                "https://www.coursera.org/projects/vuejs-intro",
                "https://www.udemy.com/course/vue-js-tutorial-free/"
            ]
        }
    },

    "Back-End Developer": {
        "skills": ["Python", "Node.js", "Java", "SQL", "REST APIs"],
        "courses": {
            "Python": [
                "https://www.coursera.org/learn/python",
                "https://www.udemy.com/course/pythonforbeginnersfree/"
            ],
            "Node.js": [
                "https://www.coursera.org/learn/nodejs-basics",
                "https://www.udemy.com/course/nodejs-tutorial-free/"
            ],
            "Java": [
                "https://www.coursera.org/learn/java-programming",
                "https://learn.microsoft.com/en-us/training/paths/java-fundamentals/"
            ],
            "SQL": [
                "https://www.coursera.org/learn/sql-for-data-science",
                "https://learn.microsoft.com/en-us/training/paths/query-relational-data-sql/"
            ],
            "REST APIs": [
                "https://www.coursera.org/learn/rest-api",
                "https://www.udemy.com/course/rest-api-tutorial-free/"
            ]
        }
    },

    "Full Stack Developer": {
        "skills": ["HTML", "CSS", "JavaScript", "React", "Node.js", "Python", "SQL"],
        "courses": {
            "HTML": ["https://www.coursera.org/learn/html-css-javascript"],
            "CSS": ["https://www.coursera.org/learn/html-css-javascript"],
            "JavaScript": ["https://www.coursera.org/learn/javascript-basics"],
            "React": ["https://www.coursera.org/learn/react-basics"],
            "Node.js": ["https://www.coursera.org/learn/nodejs-basics"],
            "Python": ["https://www.coursera.org/learn/python"],
            "SQL": ["https://www.coursera.org/learn/sql-for-data-science"]
        }
    },

    "Mobile App Developer": {
        "skills": ["Flutter", "React Native", "Swift", "Kotlin"],
        "courses": {
            "Flutter": [
                "https://www.coursera.org/projects/flutter-build-app",
                "https://www.udemy.com/course/flutter-tutorial-free/"
            ],
            "React Native": [
                "https://www.coursera.org/learn/react-native",
                "https://www.udemy.com/course/react-native-tutorial-free/"
            ],
            "Swift": [
                "https://www.coursera.org/learn/swift-programming",
                "https://www.udemy.com/course/swift-tutorial-free/"
            ],
            "Kotlin": [
                "https://www.coursera.org/learn/kotlin-for-java-developers",
                "https://www.udemy.com/course/kotlin-tutorial-free/"
            ]
        }
    },

    "Game Developer": {
        "skills": ["Unity", "Unreal Engine", "C#", "C++"],
        "courses": {
            "Unity": [
                "https://learn.unity.com/",
                "https://www.udemy.com/course/unity-tutorial-free/"
            ],
            "Unreal Engine": [
                "https://www.unrealengine.com/en-US/onlinelearning-courses",
                "https://www.udemy.com/course/unreal-engine-tutorial-free/"
            ],
            "C#": [
                "https://learn.microsoft.com/en-us/training/paths/csharp-first-steps/",
                "https://www.udemy.com/course/free-learn-csharp-tutorial-course/"
            ],
            "C++": [
                "https://www.coursera.org/specializations/c-plus-plus-modern",
                "https://www.udemy.com/course/free-learn-cpp-tutorial-course/"
            ]
        }
    },

    "UI/UX Designer": {
        "skills": ["Figma", "Adobe XD", "Sketch", "Prototyping"],
        "courses": {
            "Figma": [
                "https://www.coursera.org/learn/ui-ux-design",
                "https://www.udemy.com/course/figma-for-beginners-free/"
            ],
            "Adobe XD": ["https://www.udemy.com/course/adobe-xd-tutorial-free/"],
            "Sketch": ["https://www.udemy.com/course/sketch-tutorial-free/"],
            "Prototyping": ["https://www.coursera.org/learn/ux-design"]
        }
    },

    "Embedded Developer": {
        "skills": ["C", "C++", "Microcontrollers", "RTOS"],
        "courses": {
            "C": [
                "https://www.udemy.com/course/c-programming-for-beginners/",
                "https://www.greatlearning.in/academy/learn-for-free/courses/c-language"
            ],
            "C++": ["https://www.coursera.org/specializations/c-plus-plus-modern"],
            "Microcontrollers": ["https://www.udemy.com/course/microcontroller-tutorial-free/"],
            "RTOS": ["https://www.udemy.com/course/rtos-tutorial-free/"]
        }
    },

    # ---------------- Data / AI / ML ----------------
    "Data Analyst": {
        "skills": ["Excel", "SQL", "Python", "Tableau", "Power BI"],
        "courses": {
            "Excel": ["https://www.coursera.org/learn/excel-data-analysis"],
            "SQL": ["https://www.coursera.org/learn/sql-for-data-science"],
            "Python": ["https://www.coursera.org/learn/data-analysis-with-python"],
            "Tableau": ["https://www.coursera.org/learn/data-visualization-tableau"],
            "Power BI": ["https://www.coursera.org/learn/power-bi"]
        }
    },

    "Data Scientist": {
        "skills": ["Python", "SQL", "Pandas", "NumPy", "Scikit-learn", "Matplotlib"],
        "courses": {
            "Python": ["https://www.coursera.org/learn/python"],
            "SQL": ["https://www.coursera.org/learn/sql-for-data-science"],
            "Pandas": ["https://www.coursera.org/learn/data-analysis-with-python"],
            "NumPy": ["https://www.coursera.org/learn/numpy"],
            "Scikit-learn": ["https://www.coursera.org/learn/machine-learning"],
            "Matplotlib": ["https://www.coursera.org/learn/data-visualization-python"]
        }
    },

    "Machine Learning Engineer": {
        "skills": ["Python", "TensorFlow", "PyTorch", "Scikit-learn", "ML algorithms"],
        "courses": {
            "Python": ["https://www.coursera.org/learn/python"],
            "TensorFlow": ["https://www.coursera.org/learn/introduction-tensorflow"],
            "PyTorch": ["https://www.udemy.com/course/pytorch-tutorial-free/"],
            "Scikit-learn": ["https://www.coursera.org/learn/machine-learning"],
            "ML algorithms": ["https://www.coursera.org/learn/machine-learning"]
        }
    },

    "AI Engineer": {
        "skills": ["Python", "Deep Learning", "NLP", "TensorFlow", "PyTorch"],
        "courses": {
            "Python": ["https://www.coursera.org/learn/python"],
            "Deep Learning": ["https://www.coursera.org/specializations/deep-learning"],
            "NLP": ["https://www.coursera.org/learn/nlp-sequence-models"],
            "TensorFlow": ["https://www.coursera.org/learn/introduction-tensorflow"],
            "PyTorch": ["https://www.udemy.com/course/pytorch-tutorial-free/"]
        }
    },

    "NLP Engineer": {
        "skills": ["Python", "NLP", "SpaCy", "NLTK", "Transformers"],
        "courses": {
            "Python": ["https://www.coursera.org/learn/python"],
            "NLP": ["https://www.coursera.org/learn/language-processing"],
            "SpaCy": ["https://www.udemy.com/course/spacy-tutorial-free/"],
            "NLTK": ["https://www.udemy.com/course/nltk-tutorial-free/"],
            "Transformers": ["https://www.udemy.com/course/transformers-tutorial-free/"]
        }
    },

    "Big Data Engineer": {
        "skills": ["Hadoop", "Spark", "Python", "SQL"],
        "courses": {
            "Hadoop": ["https://www.coursera.org/learn/hadoop"],
            "Spark": ["https://www.coursera.org/learn/pyspark"],
            "Python": ["https://www.coursera.org/learn/python"],
            "SQL": ["https://www.coursera.org/learn/sql-for-data-science"]
        }
    },

    # ---------------- Cloud / DevOps / Security ----------------
    "Cloud Engineer": {
        "skills": ["AWS", "Azure", "GCP", "Docker", "Kubernetes"],
        "courses": {
            "AWS": ["https://www.coursera.org/learn/aws-cloud"],
            "Azure": ["https://www.coursera.org/learn/azure-cloud"],
            "GCP": ["https://www.coursera.org/learn/gcp-fundamentals"],
            "Docker": ["https://www.coursera.org/learn/docker"],
            "Kubernetes": ["https://www.coursera.org/learn/kubernetes-basics"]
        }
    },

    "DevOps Engineer": {
        "skills": ["Linux", "CI/CD", "Docker", "Kubernetes", "Python", "Git"],
        "courses": {
            "Linux": ["https://www.coursera.org/learn/linux"],
            "CI/CD": ["https://www.coursera.org/learn/ci-cd"],
            "Docker": ["https://www.coursera.org/learn/docker"],
            "Kubernetes": ["https://www.coursera.org/learn/kubernetes-basics"],
            "Python": ["https://www.coursera.org/learn/python"],
            "Git": ["https://www.coursera.org/learn/introduction-git-github"]
        }
    },

    "Cybersecurity Analyst": {
        "skills": ["Network Security", "Penetration Testing", "Firewalls", "Python"],
        "courses": {
            "Network Security": ["https://www.coursera.org/learn/network-security"],
            "Penetration Testing": ["https://www.coursera.org/learn/ethical-hacking"],
            "Firewalls": ["https://www.udemy.com/course/firewall-tutorial-free/"],
            "Python": ["https://www.coursera.org/learn/python"]
        }
    },

    "IT Support / Helpdesk": {
        "skills": ["Troubleshooting", "Windows OS", "Linux OS", "Networking basics"],
        "courses": {
            "Troubleshooting": ["https://www.coursera.org/professional-certificates/google-it-support"],
            "Windows OS": ["https://www.coursera.org/learn/windows"],
            "Linux OS": ["https://www.coursera.org/learn/linux"],
            "Networking basics": ["https://www.coursera.org/learn/computer-networking"]
        }
    },

    # ---------------- Database / Analytics ----------------
    "Database Administrator": {
        "skills": ["SQL", "PostgreSQL", "MySQL", "Oracle", "Performance Tuning"],
        "courses": {
            "SQL": ["https://www.coursera.org/learn/sql-for-data-science"],
            "PostgreSQL": ["https://www.coursera.org/learn/postgresql"],
            "MySQL": ["https://www.coursera.org/learn/mysql"],
            "Oracle": ["https://www.coursera.org/learn/oracle-database"],
            "Performance Tuning": ["https://www.udemy.com/course/database-performance-tuning-tutorial-free/"]
        }
    },

    "Business Analyst": {
        "skills": ["Excel", "SQL", "Data Visualization", "Power BI", "Tableau"],
        "courses": {
            "Excel": ["https://www.coursera.org/learn/excel-data-analysis"],
            "SQL": ["https://www.coursera.org/learn/sql-for-data-science"],
            "Data Visualization": ["https://www.coursera.org/learn/data-visualization-python"],
            "Power BI": ["https://www.coursera.org/learn/power-bi"],
            "Tableau": ["https://www.coursera.org/learn/data-visualization-tableau"]
        }
    },

    # ---------------- Emerging / Modern Roles ----------------
    "Robotics Engineer": {
        "skills": ["C++", "Python", "ROS", "Embedded Systems"],
        "courses": {
            "C++": ["https://www.coursera.org/specializations/c-plus-plus-modern"],
            "Python": ["https://www.coursera.org/learn/python"],
            "ROS": ["https://www.coursera.org/learn/introduction-ros"],
            "Embedded Systems": ["https://www.udemy.com/course/embedded-systems-tutorial-free/"]
        }
    },

    "Blockchain Developer": {
        "skills": ["Solidity", "Ethereum", "Smart Contracts", "Web3.js"],
        "courses": {
            "Solidity": ["https://www.udemy.com/course/solidity-tutorial-free/"],
            "Ethereum": ["https://www.coursera.org/learn/ethereum"],
            "Smart Contracts": ["https://www.coursera.org/learn/smart-contracts"],
            "Web3.js": ["https://www.udemy.com/course/web3js-tutorial-free/"]
        }
    },

    "AR/VR Developer": {
        "skills": ["Unity", "Unreal Engine", "C#", "3D Modeling"],
        "courses": {
            "Unity": ["https://learn.unity.com/"],
            "Unreal Engine": ["https://www.unrealengine.com/en-US/onlinelearning-courses"],
            "C#": ["https://learn.microsoft.com/en-us/training/paths/csharp-first-steps/"],
            "3D Modeling": ["https://www.udemy.com/course/3d-modeling-tutorial-free/"]
        }
    },

    "QA / Test Engineer": {
        "skills": ["Selenium", "JUnit", "Test Automation", "Python", "Manual Testing"],
        "courses": {
            "Selenium": ["https://www.udemy.com/course/selenium-tutorial-free/"],
            "JUnit": ["https://www.udemy.com/course/junit-tutorial-free/"],
            "Test Automation": ["https://www.udemy.com/course/automation-testing-tutorial-free/"],
            "Python": ["https://www.coursera.org/learn/python"],
            "Manual Testing": ["https://www.udemy.com/course/manual-testing-tutorial-free/"]
        }
    },

    "IoT Developer": {
        "skills": ["Arduino", "Raspberry Pi", "Python", "C", "Sensors"],
        "courses": {
            "Arduino": ["https://www.coursera.org/learn/arduino-platform"],
            "Raspberry Pi": ["https://www.coursera.org/learn/the-raspberry-pi-platform-and-python-programming-for-the-raspberry-pi"],
            "Python": ["https://www.coursera.org/learn/python"],
            "C": ["https://www.udemy.com/course/c-programming-for-beginners/"],
            "Sensors": ["https://www.coursera.org/learn/using-sensors-with-your-raspberry-pi"]
        }
    }
}

   
   



# ----------------- SKILL EXTRACTION -----------------
import re
from PyPDF2 import PdfReader
from docx import Document

def extract_skills(text):
    text = text.lower()
    keywords = []
    for role_info in role_skills_courses.values():
        for skill in role_info.get("skills", []):
            skill_lower = skill.lower()
            # For single letters like C, use word boundaries
            if skill_lower in ["c", "r"]:
                pattern = rf"\b{re.escape(skill_lower)}\b"
            else:
                # For other skills, allow punctuation after skill
                pattern = rf"\b{re.escape(skill_lower)}\b|{re.escape(skill_lower)}(?=[^\w\s])"
            if re.search(pattern, text) and skill not in keywords:
                keywords.append(skill)
    return keywords

def extract_skills_from_pdf(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + " "
        return extract_skills(text)
    except:
        return []

def extract_skills_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = " ".join([para.text for para in doc.paragraphs])
        return extract_skills(text)
    except:
        return []


def generate_word_report(profile, missing_skills, missing_courses, role_name="General Role"):
    doc = Document()
    doc.add_heading("AI Resume Analyzer Report", level=1)
    doc.add_paragraph(f"Username: {profile.get('name', 'N/A')}")
    doc.add_paragraph(f"Career Goal: {profile.get('career_goal', 'N/A')}")
    doc.add_paragraph(f"Analyzed Role: {role_name}")
    doc.add_paragraph(f"Detected Skills: {', '.join(profile.get('skills', [])) if profile.get('skills') else 'N/A'}")
    doc.add_paragraph("")

    doc.add_heading("Missing Skills & Suggested Courses", level=2)
    if missing_skills:
        for i, skill in enumerate(missing_skills):
            course = missing_courses[i] if i < len(missing_courses) else "No course available"
            doc.add_paragraph(f"{i+1}. {skill} → {course}")
    else:
        doc.add_paragraph("✅ You have all the required skills for this role!")

    # Save to an in-memory buffer instead of a temp file
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()
# ----------------- NAVIGATION -----------------
page_map = {
    "Home": "home",
    "Login": "login",
    "Register": "register",
    "Upload Resume": "resume_upload",
    "Recommendation": "recommendation",
    "Role Suggestions": "role_suggestions",
    "Skill Chart": "chart_page",
    "Feedback": "feedback",
    "Thank You": "thankyou"
}

def go_to(page_name):
    st.session_state["page"] = page_name
    st.rerun()

def sidebar_menu():
    st.sidebar.title("Navigation")
    menu_items = ["Home", "Login", "Register"]
    if st.session_state.get("user"):
        menu_items += ["Upload Resume", "Recommendation","Role Suggestions", "Skill Chart",  "Feedback", "Thank You"]

    current_page = st.session_state.get("page", "home")
    reverse_map = {v:k for k,v in page_map.items()}
    default_index = menu_items.index(reverse_map.get(current_page, "Home"))

    choice = st.sidebar.radio("Go to:", menu_items, index=default_index)
    if choice != reverse_map.get(current_page, "Home"):
        st.session_state["page"] = page_map.get(choice, "home")
        st.rerun()

# ----------------- INITIAL STATE -----------------
if "page" not in st.session_state: st.session_state["page"] = "home"
if "user" not in st.session_state: st.session_state["user"] = None
if "user_skills" not in st.session_state: st.session_state["user_skills"] = []
if "resume_path" not in st.session_state: st.session_state["resume_path"] = None
if "feedback" not in st.session_state: st.session_state["feedback"] = ""

# Render sidebar
sidebar_menu()

# ----------------- MODULES -----------------
def home_module():
    lottie_home = load_lottie("https://assets10.lottiefiles.com/packages/lf20_jcikwtux.json")
    col1, col2 = st.columns([2, 1])

    with col1:
        if lottie_home:
            st_lottie(lottie_home, height=220)
        card_container("""
            <div style='text-align: center;'>
                <h1 style='margin-bottom: 0;'>👋 Welcome</h1>
                <h1 style='
                    margin-top: 5px;
                    color: #f5f5f5;
                    font-weight: bold;
                    font-size: 42px;
                '>
                    AI Resume Analyzer
                </h1>
                <p class='small-note'>
                    Upload your resume (PDF/DOCX) to get role suggestions, missing skills and curated courses.
                </p>
            </div>
        """)

    with col2:
        st.markdown("<div class='glass-card' style='text-align:left'>", unsafe_allow_html=True)
        st.subheader("Quick Tips for Best Results ✅")
        st.write("- Supported file types: **PDF** and **DOCX** only.")
        st.write("- Highlight **projects, internships, and certifications**.")
        st.write("- Add a short **career goal statement** for personalized recommendation.")
        st.write("- Only for **computer sector students**.")
        st.markdown("</div>", unsafe_allow_html=True)


def login_module():
    card_container("<h2>🔐 Login</h2>")
    username = st.text_input("Username", key="login_username").strip()
    password = st.text_input("Password", type="password", key="login_password")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Login", key="login_btn"):
            if not username or not password:
                st.warning("Enter username & password.")
            else:
                conn = get_conn()
                if conn:
                    try:
                        cursor = conn.cursor(dictionary=True)
                        cursor.execute("SELECT * FROM a_users WHERE username=%s", (username,))
                        user = cursor.fetchone()
                        if user and check_password(password, user.get("password_hash")):
                            st.session_state["user"] = {"id": user["id"], "username": user["username"]}
                            st.success("Login successful ✅")
                            go_to("resume_upload")
                        else:
                            st.error("Invalid credentials ✖")
                    except Exception as e:
                        st.error(f"Login error: {e}")
                    finally:
                        cursor.close()
                        conn.close()
    with col2:
        if st.button("Back to Home"):
            go_to("home")

def register_module():
    card_container("<h2>📝 Create Account</h2>")
    new_user = st.text_input("Username", key="reg_username").strip().lower()
    new_email = st.text_input("Email", key="reg_email").strip().lower()
    new_pass = st.text_input("Password", type="password", key="reg_password").strip()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Register", key="reg_btn"):
            if not new_user or not new_pass or not new_email:
                st.warning("Please fill all fields including email.")
            else:
                conn = get_conn()
                if conn:
                    try:
                        cursor = conn.cursor(dictionary=True)
                        cursor.execute(
                            "SELECT id FROM a_users WHERE LOWER(username)=%s OR LOWER(email)=%s",
                            (new_user, new_email)
                        )
                        result = cursor.fetchone()
                        if result is None:
                            hashed = hash_password(new_pass)
                            cursor.execute(
                                "INSERT INTO a_users(username, email, password_hash) VALUES(%s, %s, %s)",
                                (new_user, new_email, hashed)
                            )
                            conn.commit()
                            st.success("Registration successful 🎉 — please login.")
                            go_to("login")
                        else:
                            st.error("Username or email already exists.")
                    except Exception as e:
                        st.error(f"Registration error: {e}")
                    finally:
                        cursor.close()
                        conn.close()


def resume_upload_module():
    card_container("<h2>📄 Upload Resume</h2>")
    uploaded_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
    if uploaded_file:
        save_dir = "data/sample_resumes"
        os.makedirs(save_dir, exist_ok=True)
        save_path = os.path.join(save_dir, f"{int(time.time())}_{uploaded_file.name}")
        with open(save_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.success("Resume uploaded ✅")
        st.session_state["resume_path"] = save_path

        if uploaded_file.name.lower().endswith(".pdf"):
            resume_skills = extract_skills_from_pdf(save_path)
        else:
            resume_skills = extract_skills_from_docx(save_path)

        if not resume_skills:
            st.warning("No recognized skills found. Try using a different resume or add skills manually.")
        else:
            st.session_state["user_skills"] = resume_skills
            st.success(f"Detected Skills: {', '.join(resume_skills)}")

        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Recommendations", key="resume_next_btn"):
                if resume_skills:
                    go_to("recommendation")
                else:
                    st.warning("No skills detected yet.")
        with col2:
            if st.button("Back", key="resume_back_btn"):
                go_to("home")

    else:
        st.markdown(
            '<p style="color:#ffffff;">Upload a PDF or DOCX resume to start.</p>',
            unsafe_allow_html=True
        )
        if st.button("Back", key="resume_empty_back_btn"):
            go_to("home")

# ====================== SAVE ANALYSIS FUNCTION ======================
def save_analysis_results(user_id, role, matched_skills, missing_skills, score=0):
    """Save analyzed recommendations into c_analysis_results table"""
    conn = get_conn()
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO c_analysis_results (user_id, role_name, matched_skills, missing_skills, score)
                VALUES (%s, %s, %s, %s, %s)
            """, (
                user_id,
                role,
                ",".join(matched_skills),
                ",".join(missing_skills),
                score
            ))
            conn.commit()
        except Exception as e:
            st.error(f"❌ Error saving analysis: {e}")
        finally:
            cursor.close()
            conn.close()

#==========recommendation================#
def recommendation_module():
    """Main Recommendation Page with instant navigation"""
    resume_skills = st.session_state.get("user_skills", [])
    user = st.session_state.get("user", {})
    user_id = user.get("id")

    if not resume_skills:
        st.warning("Please upload your resume first!")
        if st.button("Go to Upload Resume"):
            go_to("resume_upload")
        return

    card_container("<h2>🔎 Career Recommendations</h2>")

    # ----------------- COMPUTE ONCE -----------------
    if "recommendations_cache" not in st.session_state:
        recommendations = []

        for role, info in role_skills_courses.items():

            role_skills = info.get("skills", [])
            courses_dict = info.get("courses", {})  # <-- courses is now a dict

            matched_skills = [s for s in role_skills if s in resume_skills]

            match_score = 0

            if matched_skills:

                missing_skills = [s for s in role_skills if s not in resume_skills]

                missing_courses = []
                for skill in missing_skills:
                    if skill in courses_dict:
                        # Take the first available course link
                        course_links = courses_dict[skill]
                        if isinstance(course_links, list) and course_links:
                            missing_courses.append(course_links[0])
                        else:
                            missing_courses.append("No course available")
                    else:
                        missing_courses.append("No course available")

                if len(role_skills) > 0:
                    match_score = round(len(matched_skills) / len(role_skills) * 100, 2)

                recommendations.append({
                    "role": role,
                    "matched_skills": matched_skills,
                    "missing_skills": missing_skills,
                    "missing_courses": missing_courses,
                    "score": match_score
                })

        st.session_state["recommendations_cache"] = recommendations

    recommendations = st.session_state.get("recommendations_cache", [])

    # ----------------- DISPLAY RECOMMENDATIONS -----------------
    if not recommendations:
        st.warning("No suitable role found. Try adding more skills to your resume.")
    else:
        for rec in recommendations:
            st.markdown(f"### 🎯 {rec['role']}")
            st.markdown(f"**Match Score:** {rec['score']}%")
            st.markdown(f"- ✅ **Matched Skills:** {', '.join(rec['matched_skills'])}")

            if rec["missing_skills"]:
                st.markdown("**🧩 Missing Skills & Courses:**")
                for i, skill in enumerate(rec["missing_skills"]):
                    course = rec["missing_courses"][i] if i < len(rec["missing_courses"]) else "No course available"
                    st.markdown(f"• {skill} → [Course]({course})")

            st.markdown("---")

    # ----------------- NAVIGATION BUTTONS -----------------
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Suggestions"):
            go_to("role_suggestions")
    with col2:
        if st.button("Back"):
            go_to("resume_upload")

    # ----------------- OPTIONAL: SAVE ANALYSIS -----------------
    if st.button("💾 Save Analysis to DB"):
        if user_id and recommendations:
            try:
                for rec in recommendations:
                    save_analysis_results(
                        user_id=user_id,
                        role=rec['role'],
                        matched_skills=rec['matched_skills'],
                        missing_skills=rec['missing_skills'],
                        score=rec['score']
                    )
                st.success("✅ Analysis saved to database!")
            except Exception as e:
                st.error(f"❌ Error saving analysis: {e}")


#==============chart===============#
def chart_module():
    resume_skills = st.session_state.get("user_skills", [])
    card_container("<h2>📊 Skill Match Chart</h2>")
    roles_chart_data = []
    for role, info in role_skills_courses.items():
        role_skills = info.get("skills", [])
        matched = len([s for s in role_skills if s in resume_skills])
        missing = max(0, len(role_skills) - matched)
        if matched > 0:
            roles_chart_data.append({"role": role, "Matched": matched, "Missing": missing})
    if roles_chart_data:
        df_chart = pd.DataFrame(roles_chart_data)
        df_chart = df_chart.melt(id_vars="role", value_vars=["Matched","Missing"], var_name="Status", value_name="Count")
        fig = px.bar(df_chart, x="role", y="Count", color="Status", barmode="group", title="Skills Match Chart per Role", height=450)
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Upload resume or use demo to see chart")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Feedback"):
            go_to("feedback")
    with c2:
        if st.button("Back"):
            go_to("role_suggestions")

# ----------------- ROLE SUGGESTIONS MODULE -----------------
def role_suggestion_module():
    resume_skills = st.session_state.get("user_skills", [])
    if not resume_skills:
        st.warning("Please upload your resume first!")
        if st.button("Go to Upload Resume"):
            go_to("resume_upload")
        return

    card_container("<h2>💼 Role-Based Skill Suggestions</h2>")

    # Filter computer-related roles
    computer_roles = [role for role in role_skills_courses.keys() if "Developer" in role or "Data" in role]
    selected_role = st.selectbox("Select a Role:", computer_roles)

    role_info = role_skills_courses[selected_role]
    role_skills = role_info.get("skills", [])
    courses = role_info.get("courses", [])

    matched_skills = [s for s in role_skills if s in resume_skills]
    missing_skills = [s for s in role_skills if s not in resume_skills]

    # Build missing_courses safely
    missing_courses = []
    for skill in missing_skills:
        if skill in role_skills:
            idx = role_skills.index(skill)
            if idx < len(courses):
                missing_courses.append(courses[idx])
            else:
                missing_courses.append("No course available")
        else:
            missing_courses.append("No course available")

    st.markdown(f"### 🎯 {selected_role}")
    st.markdown(f"- ✅ **Matched Skills:** {', '.join(matched_skills) if matched_skills else 'None'}")

    if missing_skills:
        st.markdown("### ⚡ Missing Skills & Suggested Courses:")
        for skill, course in zip(missing_skills, missing_courses):
            st.markdown(f"- **{skill}** → [Course Link]({course})")
    else:
        st.success("You have all the required skills for this role! ✅")

    # Generate Word report automatically
    profile = {
        "name": st.session_state.get("user", {}).get("username", "N/A"),
        "career_goal": "Based on Resume",
        "skills": ", ".join(resume_skills)
    }
    word_bytes = generate_word_report(profile, missing_skills, missing_courses, role_name=selected_role)

    # Download button
    st.download_button(
        label="📄 Download Role Analysis Report",
        data=word_bytes,
        file_name=f"{selected_role}_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Optional chart navigation
    if st.button("📊 View Chart"):
        go_to("chart_page")



def feedback_module():
    card_container("<h2>📝 Feedback</h2>")
    feedback_text = st.text_area("Share your feedback or suggestions", value=st.session_state.get("feedback", ""))
    if st.button("Submit Feedback"):
        st.session_state["feedback"] = feedback_text
        conn = get_conn()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO b_feedback(user_id, comments) VALUES(%s, %s)",
                    (st.session_state["user"]["id"], feedback_text)
                )
                conn.commit()
                st.success("Feedback submitted! Thank you 🙏")
                go_to("thankyou")
            except Exception as e:
                st.error(f"Error saving feedback: {e}")
            finally:
                cursor.close()
                conn.close()
    if st.button("Back"):
        go_to("role_suggestions")


def thankyou_module():
    name = st.session_state.get("user", {}).get("username", "User")
    lottie_party = load_lottie("https://assets5.lottiefiles.com/packages/lf20_jbrw3hcz.json")
    if lottie_party:
        st_lottie(lottie_party, height=220)
    st.success(f"🎉 Thank you, {name}, for using AI Resume Analyzer!")
    st.markdown("You can logout below.")
    if st.button("Logout"):
        for k in list(st.session_state.keys()):
            if k != "page":
                st.session_state.pop(k)
        st.session_state["page"] = "home"        

# ----------------- PAGE ROUTING -----------------
page = st.session_state.get("page", "home")
if page == "home":
    home_module()
elif page == "login":
    login_module()
elif page == "register":
    register_module()
elif page == "resume_upload":
    resume_upload_module()
elif page == "recommendation":
    recommendation_module()
elif page == "role_suggestions":
    role_suggestion_module()    
elif page == "chart_page":
    chart_module()
elif page == "feedback":
    feedback_module()
elif page == "thankyou":
    thankyou_module()
else:
    st.error("Unknown page. Returning home.")
    st.session_state["page"] = "home"
    home_module()

